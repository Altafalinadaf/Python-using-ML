from docx import Document

doc=Document()

text="this new line adding to the existing docs file2"

doc.add_paragraph(text)
doc.add_paragraph(text)

doc.save(r'C:\Users\altaf\OneDrive\Documents\my_test.docx')
print("successfully added")